import os
import zipfile
import tempfile
from tkinter import (
    Tk, Toplevel, filedialog, Label, Button, StringVar,
    messagebox, Frame, Checkbutton, BooleanVar, Scrollbar, Canvas
)
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import openpyxl

# Fixed image size in Word
IMG_HEIGHT_IN = 1.56
IMG_WIDTH_IN = 1.04
SUPPORTED_EXTS = (".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tif", ".tiff")


def list_photos_in_zip(zip_path):
    with zipfile.ZipFile(zip_path, 'r') as z:
        members = [m for m in z.namelist() if not m.endswith("/")]
        photos = [m for m in members if os.path.splitext(m)[1].lower() in SUPPORTED_EXTS]
    return photos


def extract_selected_photos(zip_path, members, temp_dir):
    out = []
    with zipfile.ZipFile(zip_path, 'r') as z:
        for m in members:
            try:
                base = os.path.basename(m)
                raw_path = os.path.join(temp_dir, base)
                with z.open(m) as src, open(raw_path, "wb") as dst:
                    dst.write(src.read())
                img = Image.open(raw_path).convert("RGB")
                cleaned = os.path.join(temp_dir, f"cleaned_{base}.jpg")
                img.save(cleaned, "JPEG")
                display_name = os.path.splitext(base)[0]  # Last, First
                out.append((display_name, cleaned, base))
            except Exception as e:
                print(f"Skipping {m}: {e}")
    return out


def format_first_last(display_name_last_first):
    if "," in display_name_last_first:
        last, first = display_name_last_first.split(",", 1)
        return f"{first.strip()} {last.strip()}"
    return display_name_last_first


def create_word_table(photo_items, output_path):
    doc = Document()
    table = doc.add_table(rows=0, cols=6)
    table.autofit = True

    row_cells = None
    col_index = 0

    for display_name, image_path, _orig in photo_items:
        if col_index == 0:
            row_cells = table.add_row().cells

        cell = row_cells[col_index]
        p_img = cell.paragraphs[0]
        run_img = p_img.add_run()
        run_img.add_picture(image_path, width=Inches(IMG_WIDTH_IN), height=Inches(IMG_HEIGHT_IN))

        name = format_first_last(display_name)
        p_name = cell.add_paragraph(name)
        p_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_name = p_name.runs[0]
        run_name.font.size = Pt(10)
        run_name.bold = True

        col_index += 1
        if col_index == 6:
            col_index = 0

    doc.save(output_path)


def load_excel_names(excel_path):
    """Return a set of names as 'Last, First' from Excel file."""
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    sheet = wb.active
    names = set()
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if row[0] and row[1]:
            last = str(row[0]).strip()
            first = str(row[1]).strip()
            names.add(f"{last}, {first}")
    return names


class SelectionDialog(Toplevel):
    def __init__(self, parent, names_list, pre_checked=None, on_confirm=None):
        super().__init__(parent)
        self.title("Select Photos to Include")
        self.geometry("400x450")
        self.on_confirm = on_confirm

        # sort by (last, first)
        def name_sort_key(name):
            if "," in name:
                last, first = name.split(",", 1)
                return (last.strip().lower(), first.strip().lower())
            return (name.lower(), "")

        self.names = sorted(names_list, key=name_sort_key)
        self.vars = []

        container = Frame(self)
        container.pack(fill="both", expand=True, padx=8, pady=8)

        canvas = Canvas(container, borderwidth=0)
        scrollbar = Scrollbar(container, orient="vertical", command=canvas.yview)
        self.list_frame = Frame(canvas)
        self.list_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=self.list_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        pre_checked = pre_checked or set()
        for name in self.names:
            var = BooleanVar(value=name in pre_checked)
            cb = Checkbutton(self.list_frame, text=name, variable=var, anchor="w", justify="left")
            cb.pack(fill="x", anchor="w")
            self.vars.append((name, var))

        btn_frame = Frame(self)
        btn_frame.pack(fill="x", padx=8, pady=(0, 8))
        Button(btn_frame, text="Select All", command=self.select_all).pack(side="left", padx=(0, 6))
        Button(btn_frame, text="Clear All", command=self.clear_all).pack(side="left", padx=(0, 6))
        Button(btn_frame, text="Confirm", command=self.confirm).pack(side="right")
        self.grab_set()
        self.focus_set()

    def select_all(self):
        for _, var in self.vars:
            var.set(True)

    def clear_all(self):
        for _, var in self.vars:
            var.set(False)

    def confirm(self):
        selected = [name for name, var in self.vars if var.get()]
        if not selected and not messagebox.askyesno("No Selection", "No items selected. Generate empty document?"):
            return
        if self.on_confirm:
            self.on_confirm(selected)
        self.destroy()


class PhotoRosterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Photo Roster Generator")
        self.root.geometry("500x260")
        self.zip_path = None
        self.excel_path = None
        self.file_label_var = StringVar(value="Click 'Browse ZIP' to select file")

        self.drop_label = Label(
            root, textvariable=self.file_label_var,
            relief="solid", width=55, height=5, bg="#eef6ff"
        )
        self.drop_label.pack(pady=10)

        btn_row = Frame(root)
        btn_row.pack(pady=5)
        Button(btn_row, text="Browse ZIP", width=12, command=self.browse_zip).pack(side="left", padx=6)
        Button(btn_row, text="Load Excel (Optional)", width=18, command=self.browse_excel).pack(side="left", padx=6)
        Button(btn_row, text="Generate Roster", width=18, command=self.start_selection).pack(side="left", padx=6)

    def browse_zip(self):
        path = filedialog.askopenfilename(title="Select ZIP file of photos", filetypes=[("ZIP files", "*.zip")])
        if path:
            self.zip_path = path
            self.file_label_var.set(f"Selected: {os.path.basename(self.zip_path)}")

    def browse_excel(self):
        path = filedialog.askopenfilename(title="Select Excel file", filetypes=[("Excel Files", "*.xlsx *.xls")])
        if path:
            self.excel_path = path
            messagebox.showinfo("Excel Loaded", f"Loaded {os.path.basename(path)}. Only these names will be pre-checked.")

    def start_selection(self):
        if not self.zip_path:
            messagebox.showerror("Error", "Please select a ZIP file first.")
            return
        try:
            names_in_zip = list_photos_in_zip(self.zip_path)
            if not names_in_zip:
                messagebox.showerror("Error", "No valid image files found in the ZIP.")
                return

            base_names = [os.path.splitext(os.path.basename(n))[0] for n in names_in_zip]
            pre_checked = set()
            if self.excel_path:
                pre_checked = load_excel_names(self.excel_path)

            def on_confirm(selected_display_names):
                selected_members = []
                for member in names_in_zip:
                    base = os.path.splitext(os.path.basename(member))[0]
                    if base in selected_display_names:
                        selected_members.append(member)
                self.generate_roster(selected_members)

            SelectionDialog(self.root, base_names, pre_checked=pre_checked, on_confirm=on_confirm)

        except Exception as e:
            messagebox.showerror("Error", f"Failed to read ZIP: {e}")

    def generate_roster(self, selected_members):
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                photos = extract_selected_photos(self.zip_path, selected_members, temp_dir)

                # sort photos by last name, then first name
                def photo_sort_key(item):
                    display_name, _, _ = item
                    if "," in display_name:
                        last, first = display_name.split(",", 1)
                        return (last.strip().lower(), first.strip().lower())
                    return (display_name.lower(), "")

                photos.sort(key=photo_sort_key)

                if not photos:
                    messagebox.showerror("Error", "No valid photos after extraction.")
                    return
                output_path = os.path.join(os.path.dirname(self.zip_path), "PhotoRoster.docx")
                create_word_table(photos, output_path)
                messagebox.showinfo("Success", f"Word roster created:\n{output_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create roster: {e}")


def main():
    root = Tk()
    app = PhotoRosterApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
