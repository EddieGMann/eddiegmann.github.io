import tkinter as tk
from tkinter import filedialog, messagebox
from tkinterdnd2 import TkinterDnD, DND_FILES
import csv
from collections import defaultdict
from docx import Document
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import statistics
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.shared import OxmlElement as OxmlElt

# Page width for proportional columns
PAGE_WIDTH_CM = 16  # adjust as needed
FIRST_ROW_HEIGHT_CM = 2.75  # height of first row in cm
LIGHT_GRAY_HEX = "D3D3D3"

# Label sets
SATISFIED_SET = ["Very Dissatisfied", "Dissatisfied", "Neutral", "Somewhat Satisfied", "Very Satisfied"]
AGREE_SET     = ["Strongly Disagree", "Disagree", "Neutral", "Agree", "Strongly Agree"]
CONFIDENT_SET = ["No Confidence", "Limited Confidence", "Somewhat Confident", "Confident", "Very Confident"]
QUALITYTA_SET   = ["Poor", "Fair", "Good","Excellent"]
QUALITYPROF_SET = ["Poor", "Fair", "Satisfactory", "Good", "Excellent"]
IMPROVEMENT_SET = ["Not At All", "Limited Improvement", "Noticeable Improvement", "Significant Improvement", "Outstanding Improvement"]

CANDIDATE_SETS = [
    ("SATISFIED", SATISFIED_SET),
    ("AGREE", AGREE_SET),
    ("CONFIDENT", CONFIDENT_SET),
    ("QUALITYTA", QUALITYTA_SET),
    ("QUALITYPROF", QUALITYPROF_SET),
    ("IMPROVEMENT", IMPROVEMENT_SET),
]

def set_text_direction(cell, direction="tbRl"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)
    tcPr.append(textDirection)

def clean_answer(answer: str) -> str:
    if not answer:
        return ""
    return re.sub(r"^\s*\d+\s*-\s*", "", str(answer)).strip()

def extract_main_question(q: str) -> str:
    if not q:
        return ""
    q = str(q)
    protected = re.sub(r'\bDr\.', 'Dr<DOT>', q, flags=re.IGNORECASE)
    idxs = [i for i in [protected.find("."), protected.find(":"), protected.find("?")] if i != -1]
    if idxs:
        protected = protected[:min(idxs)].strip()
    else:
        protected = protected.strip()
    return protected.replace('Dr<DOT>', 'Dr.')

def extract_sub_question(q: str) -> str:
    if not q:
        return ""
    q = str(q).strip()
    lower_q = q.lower()
    if "on a scale" in lower_q:
        parts = re.split(r"\s*-\s*", q)
        if len(parts) >= 3:
            return parts[2].strip()
        elif len(parts) >= 2:
            return parts[1].strip()
        else:
            return q
    else:
        if "-" in q:
            return q.rsplit("-", 1)[-1].strip()
        return q

def detect_label_set(answers):
    cleaned = [clean_answer(a).lower() for a in answers if a]
    if not cleaned:
        return None
    best_set = None
    best_count = 0
    for _, label_set in CANDIDATE_SETS:
        labels_lower = [s.lower() for s in label_set]
        count = sum(1 for a in cleaned if a in labels_lower)
        if count > best_count:
            best_set = label_set
            best_count = count
    return best_set if best_count > 0 else None

def is_numeric_ranking(answers):
    cleaned = [clean_answer(a) for a in answers if a]
    return all(re.fullmatch(r"\d+", a) for a in cleaned)

def compute_median(counts, answer_set):
    expanded = []
    for idx, label in enumerate(answer_set, start=1):
        expanded.extend([idx] * counts.get(label, 0))
    if expanded:
        return statistics.median(expanded)
    return ""

def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)

def apply_proportional_widths(table):
    n_cols = len(table.columns)
    first_col_width = PAGE_WIDTH_CM * 0.55
    other_col_width = (PAGE_WIDTH_CM - first_col_width) / max(1, n_cols - 1)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if col_idx == 0:
                set_cell_width(cell, first_col_width)
            else:
                set_cell_width(cell, other_col_width)
        if row_idx == 0:
            row.height = Cm(FIRST_ROW_HEIGHT_CM)
            row.height_rule = True

def handle_file_drop(event):
    file_path = event.data.strip().strip("{}")
    if file_path.lower().endswith('.csv'):
        label_status.config(text=f"File loaded: {file_path}", fg="green")
        try:
            generate_tables_based_on_prefix(file_path)
        except Exception as e:
            messagebox.showerror("Error", str(e))
    else:
        label_status.config(text="Please drop a valid CSV file.", fg="red")

def browse_file():
    file_path = filedialog.askopenfilename(filetypes=[("CSV Files", "*.csv")])
    if file_path:
        try:
            generate_tables_based_on_prefix(file_path)
        except Exception as e:
            messagebox.showerror("Error", str(e))

def set_table_title(doc, text):
    if not text:
        text = "   "
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

def color_cell(cell, hex_color):
    shading_elm = OxmlElt('w:shd')
    shading_elm.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_tables_based_on_prefix(file_path):
    with open(file_path, newline='', encoding='utf-8') as csvfile:
        reader = csv.reader(csvfile)
        data = list(reader)

    if len(data) < 4 or len(data[0]) <= 17:
        raise ValueError("CSV file does not contain enough data/columns.")

    processed_data = [row[17:] for row in data]
    header_row = processed_data[0]
    question_row = processed_data[1]
    rows = processed_data[3:]

    grouped_columns = defaultdict(list)
    for idx, header in enumerate(header_row):
        if header:
            prefix = header.split("_")[0].lower()
            grouped_columns[prefix].append(idx)

    doc = Document()
    doc.add_heading("CSV Data Grouped by Prefix", level=1)

    for prefix, column_indices in grouped_columns.items():
        # Determine title
        first_q_text = str(question_row[column_indices[0]]).lower() if column_indices else ""
        title_text = ""
        if "please rate the instructor for this course" in first_q_text:
            title_text = "Instructor Ratings Median and Frequency Responses"
        elif "please rate the teaching assistant for this course" in first_q_text:
            title_text = "TA Ratings Median and Frequency of Responses"
        elif prefix == "q1":
            title_text = "General Course Ratings Median and Frequency Responses"
        elif prefix == "q3":
            title_text = "Course Format Ratings Median and Frequency of Responses"

        set_table_title(doc, title_text)

        is_rating_block = any("_" in header_row[i] or "rate" in str(question_row[i]).lower() for i in column_indices)

        if is_rating_block:
            all_answers = []
            for i in column_indices:
                for r in rows:
                    if len(r) > i and r[i]:
                        all_answers.append(r[i])

            if is_numeric_ranking(all_answers):
                numbers_list = [int(clean_answer(a)) for a in all_answers if clean_answer(a).isdigit()]
                max_num = max(numbers_list) if numbers_list else 0
                min_num = min(numbers_list) if numbers_list else 1

                base_q = extract_main_question(question_row[column_indices[0]])
                table = doc.add_table(rows=len(column_indices) + 1, cols=max_num + 2)
                table.style = "Table Grid"

                table.rows[0].cells[0].text = base_q
                table.rows[0].cells[1].text = "Median"
                set_text_direction(table.rows[0].cells[1])

                numbers = list(range(max_num, min_num - 1, -1))
                for col_idx, num in enumerate(numbers, start=2):
                    if num == max_num:
                        label = f"{num} Least Effective"
                    elif num == min_num:
                        label = f"{num} Most Effective"
                    else:
                        label = str(num)
                    table.rows[0].cells[col_idx].text = label
                    set_text_direction(table.rows[0].cells[col_idx])

                for row_idx, col_index in enumerate(column_indices, start=1):
                    row_label = extract_sub_question(question_row[col_index])
                    table.rows[row_idx].cells[0].text = row_label

                    counts = {n: 0 for n in range(1, max_num + 1)}
                    for r in rows:
                        if len(r) > col_index and r[col_index]:
                            v = clean_answer(r[col_index])
                            if v.isdigit():
                                counts[int(v)] += 1

                    expanded = []
                    for idx in range(1, max_num + 1):
                        expanded.extend([idx] * counts[idx])
                    median_val = statistics.median(expanded) if expanded else ""
                    table.rows[row_idx].cells[1].text = str(median_val)

                    # Color median column
                    color_cell(table.rows[row_idx].cells[1], LIGHT_GRAY_HEX)

                    # Fill counts and alternate row shading including row header
                    reversed_numbers = list(range(max_num, min_num - 1, -1))
                    for col_idx, num in enumerate(reversed_numbers, start=2):
                        table.rows[row_idx].cells[col_idx].text = str(counts[num])

                    if row_idx % 2 == 1:
                        for cell in table.rows[row_idx].cells:
                            color_cell(cell, LIGHT_GRAY_HEX)

                # Color header row median cell
                color_cell(table.rows[0].cells[1], LIGHT_GRAY_HEX)

                apply_proportional_widths(table)

            else:
                use_labels = detect_label_set(all_answers)
                if not use_labels:
                    continue
                base_q = extract_main_question(question_row[column_indices[0]])
                table = doc.add_table(rows=len(column_indices) + 1, cols=len(use_labels) + 2)
                table.style = "Table Grid"

                table.rows[0].cells[0].text = base_q
                table.rows[0].cells[1].text = "Median"
                set_text_direction(table.rows[0].cells[1])
                for col_idx, label in enumerate(use_labels, start=2):
                    table.rows[0].cells[col_idx].text = label
                    set_text_direction(table.rows[0].cells[col_idx])

                labels_lower = [lbl.lower() for lbl in use_labels]
                for row_idx, col_index in enumerate(column_indices, start=1):
                    row_label = extract_sub_question(question_row[col_index])
                    table.rows[row_idx].cells[0].text = row_label

                    counts = {lab: 0 for lab in use_labels}
                    for r in rows:
                        if len(r) > col_index and r[col_index]:
                            ans = clean_answer(r[col_index]).lower()
                            for lab, lab_low in zip(use_labels, labels_lower):
                                if ans == lab_low:
                                    counts[lab] += 1
                                    break

                    median_val = compute_median(counts, use_labels)
                    table.rows[row_idx].cells[1].text = str(median_val)

                    # Color median column
                    color_cell(table.rows[row_idx].cells[1], LIGHT_GRAY_HEX)

                    for col_idx, lab in enumerate(use_labels, start=2):
                        table.rows[row_idx].cells[col_idx].text = str(counts[lab])

                    # Alternate row shading including row header
                    if row_idx % 2 == 1:
                        for cell in table.rows[row_idx].cells:
                            color_cell(cell, LIGHT_GRAY_HEX)

                # Color header row median cell
                color_cell(table.rows[0].cells[1], LIGHT_GRAY_HEX)

                apply_proportional_widths(table)

        else:
            for col_index in column_indices:
                column_data = [(i, r[col_index]) for i, r in enumerate(rows) if len(r) > col_index and r[col_index]]
                if not column_data:
                    column_data = [("N/A", "None Submitted")]

                title = extract_main_question(question_row[col_index]) if col_index < len(question_row) else "Unknown Question"
                set_table_title(doc, title)

                table = doc.add_table(rows=len(column_data), cols=1)
                table.style = "Table Grid"

                for row_idx, (orig_i, cell_val) in enumerate(column_data):
                    if orig_i == "N/A":
                        table.rows[row_idx].cells[0].text = cell_val
                    else:
                        student_number = orig_i + 1
                        table.rows[row_idx].cells[0].text = f"Student {student_number}: {cell_val}"

                    # Alternate row shading
                    if row_idx % 2 == 1:
                        color_cell(table.rows[row_idx].cells[0], LIGHT_GRAY_HEX)

    save_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")],
        title="Save Word Document"
    )
    if save_path:
        doc.save(save_path)
        messagebox.showinfo("Success", f"Word document created successfully at:\n{save_path}")
    else:
        messagebox.showwarning("Warning", "No save location selected.")

# GUI setup
window = TkinterDnD.Tk()
window.title("CSV to Word Converter")
window.geometry("700x400")
window.config(bg="#f0f0f0")

label_instruction = tk.Label(window, text="Drag and drop a CSV file below or click Browse.",
                             font=("Arial", 14), bg="#f0f0f0")
label_instruction.pack(pady=10)

drop_area = tk.Label(window, text="Drop CSV File Here", bg="#d9e8f5", fg="#00509e",
                     width=40, height=5, relief="ridge", font=("Arial", 12, "bold"))
drop_area.pack(pady=20)
drop_area.drop_target_register(DND_FILES)
drop_area.dnd_bind('<<Drop>>', handle_file_drop)

label_status = tk.Label(window, text="", font=("Arial", 12), bg="#f0f0f0", fg="#00509e")
label_status.pack(pady=5)

btn_browse = tk.Button(window, text="Browse", command=browse_file,
                       font=("Arial", 12), bg="#00509e", fg="white", width=15)
btn_browse.pack(pady=10)

window.mainloop()
